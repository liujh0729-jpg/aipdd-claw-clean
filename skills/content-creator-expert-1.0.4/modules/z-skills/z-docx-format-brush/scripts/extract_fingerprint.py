#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""提取 docx 格式指纹：样式表定义 + 段落/run 实测分布 + 表格/封面/分页情况。

用法:
    python3 extract_fingerprint.py 模板.docx [--json fingerprint.json]

实测优先原则：中文公文 docx 常在 run 层覆盖样式，样式表不可全信，
所以正文/标题参数均按 run 实测众数输出。
"""
import argparse
import json
import sys
from collections import Counter

import docx
from docx.oxml.ns import qn


def pt(size):
    return round(size / 12700, 1) if size else None


def run_eastasia(r):
    rPr = r.font.element.rPr
    if rPr is not None and rPr.rFonts is not None:
        return rPr.rFonts.get(qn('w:eastAsia'))
    return None


def run_color(r):
    try:
        if r.font.color and r.font.color.type is not None:
            return str(r.font.color.rgb)
    except Exception:
        pass
    return None


def style_info(doc, name):
    try:
        s = doc.styles[name]
    except KeyError:
        return None
    ea = None
    rPr = s.element.rPr
    if rPr is not None and rPr.rFonts is not None:
        ea = rPr.rFonts.get(qn('w:eastAsia'))
    color = None
    try:
        if s.font.color and s.font.color.type is not None:
            color = str(s.font.color.rgb)
    except Exception:
        pass
    return {
        'font': s.font.name, 'eastAsia': ea,
        'size_pt': pt(s.font.size.emu) if s.font.size else None,
        'bold': s.font.bold, 'color': color,
    }


def most_common(counter):
    return counter.most_common(1)[0][0] if counter else None


def extract(path):
    doc = docx.Document(path)
    fp = {'file': path, 'styles': {}, 'body': {}, 'headings': {}, 'cover': [], 'tables': {}, 'chapters': {}}

    for sn in ['Normal', 'Body Text', 'First Paragraph', 'List Paragraph',
               'Heading 1', 'Heading 2', 'Heading 3', 'Normal (Web)']:
        info = style_info(doc, sn)
        if info:
            fp['styles'][sn] = info

    # 封面：前3个非空段落
    nonempty = [p for p in doc.paragraphs if p.text.strip()]
    for p in nonempty[:3]:
        runs = [{'text': r.text[:20], 'font': r.font.name, 'eastAsia': run_eastasia(r),
                 'size_pt': pt(r.font.size.emu) if r.font.size else None, 'bold': r.bold}
                for r in p.runs[:3]]
        fp['cover'].append({'style': p.style.name, 'align': str(p.alignment), 'runs': runs})

    # 正文/标题 run 实测（跳过目录 toc 样式）
    body_font, body_size, body_indent, body_spacing = Counter(), Counter(), Counter(), Counter()
    head = {}
    page_break_chapters = 0
    total_chapters = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or p.style.name.startswith('toc'):
            continue
        st = p.style.name
        pf = p.paragraph_format
        if st.startswith('Heading'):
            hc = head.setdefault(st, {'font': Counter(), 'size': Counter(), 'bold': Counter(), 'align': Counter()})
            hc['align'][str(p.alignment)] += 1
            for r in p.runs:
                hc['font'][(r.font.name, run_eastasia(r))] += 1
                hc['size'][pt(r.font.size.emu) if r.font.size else None] += 1
                hc['bold'][bool(r.bold)] += 1
            if st == 'Heading 1' and t.startswith('第') and '章' in t[:6]:
                total_chapters += 1
                pPr = p._p.pPr
                if pPr is not None and pPr.find(qn('w:pageBreakBefore')) is not None:
                    page_break_chapters += 1
        else:
            body_indent[pf.first_line_indent.emu if pf.first_line_indent else None] += 1
            body_spacing[str(pf.line_spacing)] += 1
            for r in p.runs:
                body_font[(r.font.name, run_eastasia(r))] += 1
                body_size[pt(r.font.size.emu) if r.font.size else None] += 1

    fp['body'] = {
        'font': most_common(body_font),
        'size_pt': most_common(body_size),
        'first_line_indent_emu': most_common(body_indent),
        'line_spacing': most_common(body_spacing),
        'font_dist': [(str(k), v) for k, v in body_font.most_common(5)],
        'size_dist': [(str(k), v) for k, v in body_size.most_common(5)],
    }
    for st, hc in head.items():
        fp['headings'][st] = {
            'font': most_common(hc['font']),
            'size_pt': most_common(hc['size']),
            'bold': most_common(hc['bold']),
            'align': most_common(hc['align']),
        }
    fp['chapters'] = {'total': total_chapters, 'with_page_break': page_break_chapters}

    # 表格
    cell_font, cell_size = Counter(), Counter()
    bordered = 0
    for tb in doc.tables:
        if tb._tbl.tblPr is not None and tb._tbl.tblPr.find(qn('w:tblBorders')) is not None:
            bordered += 1
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        cell_font[(r.font.name, run_eastasia(r))] += 1
                        cell_size[pt(r.font.size.emu) if r.font.size else None] += 1
    fp['tables'] = {
        'count': len(doc.tables), 'with_borders': bordered,
        'cell_font': most_common(cell_font), 'cell_size_pt': most_common(cell_size),
    }

    # run 颜色分布（诊断 pandoc 蓝标题等）
    colors = Counter()
    for p in doc.paragraphs:
        for r in p.runs:
            c = run_color(r)
            if c:
                colors[c] += 1
    fp['run_colors'] = dict(colors.most_common(5))
    return fp


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx_path')
    ap.add_argument('--json', help='输出 JSON 指纹文件路径')
    args = ap.parse_args()

    fp = extract(args.docx_path)

    def jsafe(o):
        if isinstance(o, tuple):
            return list(o)
        return str(o)

    print('=' * 60)
    print('格式指纹:', args.docx_path)
    print('=' * 60)
    print('\n[样式表定义]')
    for sn, info in fp['styles'].items():
        print(f"  {sn}: {info}")
    print('\n[正文实测(众数)]')
    print(' ', fp['body'])
    print('\n[标题实测(众数)]')
    for st, info in fp['headings'].items():
        print(f"  {st}: {info}")
    print('\n[封面前3段]')
    for c in fp['cover']:
        print(' ', c)
    print('\n[表格]', fp['tables'])
    print('[章节分页]', fp['chapters'])
    print('[run颜色]', fp['run_colors'])

    if args.json:
        with open(args.json, 'w', encoding='utf-8') as f:
            json.dump(fp, f, ensure_ascii=False, indent=2, default=jsafe)
        print('\nJSON 已写入:', args.json)


if __name__ == '__main__':
    sys.exit(main())
