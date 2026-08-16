#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""格式统一验证：检查 run 字体收敛度、遗留颜色、表格边框、章节分页、引号方向。

用法:
    python3 verify_format.py 输出.docx [--expect-font 仿宋]
退出码 0=通过，1=有问题。
"""
import argparse
import sys
from collections import Counter

import docx
from docx.oxml.ns import qn


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx_path')
    ap.add_argument('--expect-font', default='仿宋')
    args = ap.parse_args()

    d = docx.Document(args.docx_path)
    problems = []

    # 1) run 字体/字号收敛度
    fonts = Counter()
    colors = Counter()
    for p in d.paragraphs:
        if not p.text.strip() or p.style.name.startswith('toc'):
            continue
        for r in p.runs:
            if not r.text.strip():  # 纯换行/空 run 不参与统计
                continue
            ea = None
            rPr = r.font.element.rPr
            if rPr is not None and rPr.rFonts is not None:
                ea = rPr.rFonts.get(qn('w:eastAsia'))
            sz = round(r.font.size.emu / 12700, 1) if r.font.size else None
            fonts[(r.font.name, ea, sz)] += 1
            try:
                if r.font.color and r.font.color.type is not None:
                    colors[str(r.font.color.rgb)] += 1
            except Exception:
                pass
    print('[run 字体分布]')
    for k, v in fonts.most_common(10):
        print(' ', k, v)
    bad = [(k, v) for k, v in fonts.items()
           if k[0] != args.expect_font or k[1] != args.expect_font or k[2] is None]
    if bad:
        problems.append(f'存在未统一的 run: {bad[:5]}')
    if colors:
        problems.append(f'遗留颜色: {dict(colors.most_common(3))}')

    # 2) 表格边框
    no_border = sum(1 for tb in d.tables
                    if tb._tbl.tblPr is None or tb._tbl.tblPr.find(qn('w:tblBorders')) is None)
    print(f'[表格] 共{len(d.tables)}个，无边框{no_border}个')
    if no_border:
        problems.append(f'{no_border}个表格无边框')

    # 3) 章节分页
    for p in d.paragraphs:
        t = p.text.strip()
        if p.style.name == 'Heading 1' and t.startswith('第') and '章' in t[:6]:
            pPr = p._p.pPr
            has = pPr is not None and pPr.find(qn('w:pageBreakBefore')) is not None
            print(('分页 ' if has else '无分页'), '|', t[:24])
            if not has:
                problems.append(f'章节未分页: {t[:20]}')

    # 4) 引号方向抽查
    import re
    for p in d.paragraphs:
        for m in re.finditer(r'[“”"]', p.text):
            pass
        if '”' in p.text and '“' in p.text:
            if p.text.index('”') < p.text.index('“'):
                problems.append(f'引号方向错误: {p.text[:40]}')
                break

    print()
    if problems:
        print('未通过:')
        for x in problems:
            print(' -', x)
        return 1
    print('全部通过：字体收敛、无杂色、表格有边框、章节分页、引号正确')
    return 0


if __name__ == '__main__':
    sys.exit(main())
