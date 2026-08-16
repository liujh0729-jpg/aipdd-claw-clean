#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""根据格式指纹和结构化内容生成新的 docx。

用法:
    python3 gen_from_template.py content.json --config fingerprint.json --out output.docx

content.json 示例:
{
  "cover": {
    "title": ["示例公司", "示例项目"],
    "document_type": "报告",
    "metadata": ["编制单位：示例公司", "日期：2026年8月"]
  },
  "blocks": [
    {"type": "heading", "level": 1, "text": "第一章 项目概述"},
    {"type": "paragraph", "text": "这里填写正文。"},
    {
      "type": "table",
      "headers": ["字段", "内容"],
      "rows": [["项目", "示例项目"], ["状态", "进行中"]]
    }
  ]
}

格式配置兼容 extract_fingerprint.py --json 的产物，也支持 apply_format.py
文档注释中的精简配置。所有标题、正文和表格均通过统一工厂函数写入，
避免逐段设置造成格式漂移。
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Pt

from apply_format import (
    add_page_break_before,
    fix_quotes_text,
    fix_style,
    load_config,
    make_set_run,
    set_table_borders,
)


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _text(value: Any) -> str:
    return "" if value is None else str(value)


def _alignment(value: str | None, default=WD_ALIGN_PARAGRAPH.LEFT):
    if value is None:
        return default
    try:
        return ALIGNMENTS[value.lower()]
    except (AttributeError, KeyError) as exc:
        choices = ", ".join(ALIGNMENTS)
        raise ValueError(f"不支持的对齐方式 {value!r}，可选值：{choices}") from exc


def configure_document(doc, cfg: dict[str, Any]) -> None:
    """一次性设置样式兜底，工厂函数仍会给每个 run 写入明确格式。"""
    font_name = cfg["font"]
    body = cfg["body"]
    for style_name in (
        "Normal",
        "Body Text",
        "First Paragraph",
        "List Paragraph",
        "Normal (Web)",
    ):
        fix_style(doc, style_name, font_name, body["size_pt"], None)
    for style_name, info in cfg["headings"].items():
        fix_style(doc, style_name, font_name, info["size_pt"], True)


def title(doc, text: str, level: int, cfg: dict[str, Any]):
    """写入统一格式的 Heading 1-3 标题。"""
    if level not in (1, 2, 3):
        raise ValueError("标题 level 仅支持 1、2、3")
    style_name = f"Heading {level}"
    size = cfg["headings"].get(style_name, {}).get(
        "size_pt", cfg["body"]["size_pt"]
    )
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = cfg["body"]["line_spacing"]
    run = p.add_run(fix_quotes_text(_text(text)))
    make_set_run(cfg["font"])(run, size, True)
    if (
        level == 1
        and cfg.get("chapter_page_break", True)
        and text.startswith("第")
        and "章" in text[:8]
    ):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_break_before(p)
    return p


def para(
    doc,
    text: str,
    cfg: dict[str, Any],
    *,
    bold: bool = False,
    align: str | None = None,
    indent: bool = True,
):
    """写入统一格式的正文段落。"""
    p = doc.add_paragraph(style="Normal")
    p.alignment = _alignment(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.line_spacing = cfg["body"]["line_spacing"]
    p.paragraph_format.first_line_indent = (
        Emu(cfg["body"]["first_line_indent_emu"]) if indent else None
    )
    run = p.add_run(fix_quotes_text(_text(text)))
    make_set_run(cfg["font"])(run, cfg["body"]["size_pt"], bold or None)
    return p


def cover(doc, data: dict[str, Any], cfg: dict[str, Any]) -> None:
    """生成标题、文档类型和元数据均显式格式化的封面。"""
    raw_title = data.get("title", [])
    title_lines = raw_title if isinstance(raw_title, list) else [raw_title]
    title_lines = [_text(x).strip() for x in title_lines if _text(x).strip()]
    document_type = _text(data.get("document_type", "")).strip()
    raw_metadata = data.get("metadata", [])
    metadata = raw_metadata if isinstance(raw_metadata, list) else [raw_metadata]

    for index, line in enumerate(title_lines):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_before = Pt(72 if index == 0 else 0)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(fix_quotes_text(line))
        make_set_run(cfg["font"])(run, cfg["cover_title_pt"], True)

    if document_type:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(72)
        run = p.add_run(fix_quotes_text(document_type))
        make_set_run(cfg["font"])(run, cfg["cover_doctype_pt"], True)

    for line in metadata:
        if not _text(line).strip():
            continue
        para(doc, _text(line), cfg, bold=True, align="center", indent=False)

    if title_lines or document_type or metadata:
        doc.add_page_break()


def table(doc, headers: list[Any], rows: list[list[Any]], cfg: dict[str, Any]):
    """写入带单线边框、统一字体和加粗表头的表格。"""
    if not isinstance(headers, list) or not headers:
        raise ValueError("table.headers 必须是非空数组")
    if not isinstance(rows, list):
        raise ValueError("table.rows 必须是数组")
    width = len(headers)
    for index, row in enumerate(rows, start=1):
        if not isinstance(row, list) or len(row) != width:
            raise ValueError(f"table.rows 第 {index} 行列数应为 {width}")

    tb = doc.add_table(rows=1, cols=width)
    tb.style = "Table Grid"
    set_table_borders(tb)
    set_run = make_set_run(cfg["font"])

    for index, value in enumerate(headers):
        cell = tb.rows[0].cells[index]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(fix_quotes_text(_text(value)))
        set_run(run, cfg["table_cell_size_pt"], True)

    for row_data in rows:
        cells = tb.add_row().cells
        for index, value in enumerate(row_data):
            cell = cells[index]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(fix_quotes_text(_text(value)))
            set_run(run, cfg["table_cell_size_pt"], None)
    return tb


def render_block(doc, block: dict[str, Any], cfg: dict[str, Any]) -> None:
    if not isinstance(block, dict):
        raise ValueError("blocks 中的每一项都必须是对象")
    block_type = block.get("type")
    if block_type == "heading":
        title(doc, _text(block.get("text")), int(block.get("level", 1)), cfg)
    elif block_type == "paragraph":
        para(
            doc,
            _text(block.get("text")),
            cfg,
            bold=bool(block.get("bold", False)),
            align=block.get("align"),
            indent=bool(block.get("indent", True)),
        )
    elif block_type == "table":
        table(doc, block.get("headers"), block.get("rows", []), cfg)
    elif block_type == "page_break":
        doc.add_page_break()
    else:
        raise ValueError(f"不支持的 block.type：{block_type!r}")


def build_document(
    content: dict[str, Any], cfg: dict[str, Any], out_path: str | Path
) -> Path:
    if not isinstance(content, dict):
        raise ValueError("内容 JSON 顶层必须是对象")
    blocks = content.get("blocks", [])
    if not isinstance(blocks, list):
        raise ValueError("content.blocks 必须是数组")

    doc = Document()
    configure_document(doc, cfg)
    if content.get("cover"):
        if not isinstance(content["cover"], dict):
            raise ValueError("content.cover 必须是对象")
        cover(doc, content["cover"], cfg)
    for block in blocks:
        render_block(doc, block, cfg)

    destination = Path(out_path).expanduser()
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)
    return destination


def main() -> int:
    parser = argparse.ArgumentParser(description="根据格式指纹和结构化内容生成 docx")
    parser.add_argument("content_json", help="结构化内容 JSON")
    parser.add_argument("--config", help="extract_fingerprint.py 生成的格式指纹 JSON")
    parser.add_argument("--out", required=True, help="输出 docx 路径")
    args = parser.parse_args()

    try:
        with open(args.content_json, encoding="utf-8") as handle:
            content = json.load(handle)
        cfg = load_config(args.config)
        destination = build_document(content, cfg, args.out)
    except (OSError, ValueError, TypeError, json.JSONDecodeError) as exc:
        print(f"生成失败：{exc}", file=sys.stderr)
        return 2

    print(f"done: {destination}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
