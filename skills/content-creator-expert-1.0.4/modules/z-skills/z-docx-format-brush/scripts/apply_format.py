#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""docx 格式刷：按格式配置统一目标文档的格式（五层：样式/段落/封面/表格/文字）。

用法:
    python3 apply_format.py 目标.docx [--config fingerprint.json] [--out 输出.docx]
    不带 --out 则原地保存；不带 --config 用中文公文默认参数（仿宋12磅正文/15磅标题）。

配置可用 extract_fingerprint.py --json 从模板生成，也可手写精简版:
    {"body": {"size_pt": 12, "first_line_indent_emu": 304800, "line_spacing": 1.5},
     "headings": {"Heading 1": {"size_pt": 15}, "Heading 2": {"size_pt": 15}, "Heading 3": {"size_pt": 12}},
     "font": "仿宋", "table_cell_size_pt": 12, "cover": true, "chapter_page_break": true}
"""
import argparse
import copy
import json
import sys

from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEFAULTS = {
    'font': '仿宋',
    'body': {'size_pt': 12, 'first_line_indent_emu': 304800, 'line_spacing': 1.5},
    'headings': {'Heading 1': {'size_pt': 15}, 'Heading 2': {'size_pt': 15}, 'Heading 3': {'size_pt': 12}},
    'table_cell_size_pt': 12,
    'cover': True,
    'cover_title_pt': 22,
    'cover_doctype_pt': 26,
    'chapter_page_break': True,
    'fix_quotes': True,
}


def load_config(path):
    # 每次加载都复制嵌套配置，避免一次自定义配置污染后续调用。
    cfg = copy.deepcopy(DEFAULTS)
    if not path:
        return cfg
    with open(path, encoding='utf-8') as f:
        raw = json.load(f)
    # 兼容 extract_fingerprint.py 的完整指纹结构
    if 'body' in raw and isinstance(raw['body'], dict):
        b = raw['body']
        cfg['body'] = {
            'size_pt': float(b.get('size_pt') or 12),
            'first_line_indent_emu': int(b['first_line_indent_emu']) if b.get('first_line_indent_emu') not in (None, 'None') else 304800,
            'line_spacing': float(b['line_spacing']) if str(b.get('line_spacing', '1.5')).replace('.', '', 1).isdigit() else 1.5,
        }
        fnt = b.get('font')
        if isinstance(fnt, (list, tuple)) and fnt[0]:
            cfg['font'] = fnt[0]
    if 'headings' in raw:
        for st, info in raw['headings'].items():
            if info.get('size_pt'):
                cfg['headings'].setdefault(st, {})['size_pt'] = float(info['size_pt'])
    if 'tables' in raw and raw['tables'].get('cell_size_pt'):
        cfg['table_cell_size_pt'] = float(raw['tables']['cell_size_pt'])
    # 手写精简版直接覆盖
    for k in ('font', 'table_cell_size_pt', 'cover', 'chapter_page_break', 'fix_quotes',
              'cover_title_pt', 'cover_doctype_pt'):
        if k in raw:
            cfg[k] = raw[k]
    return cfg


def make_set_run(font_name):
    def set_run(r, size, bold=None):
        r.font.name = font_name
        rPr = r.font.element.get_or_add_rPr()
        rPr.get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
        r.font.size = Pt(size)
        if bold is not None:
            r.bold = bold
        for c in rPr.findall(qn('w:color')):  # 清除颜色（如pandoc主题蓝）
            rPr.remove(c)
    return set_run


def fix_style(doc, sname, font_name, size, bold):
    try:
        s = doc.styles[sname]
    except KeyError:
        return
    s.font.name = font_name
    s.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
    s.font.size = Pt(size)
    s.font.bold = bold
    rPr = s.element.rPr
    if rPr is not None:
        for c in rPr.findall(qn('w:color')):
            rPr.remove(c)


def add_page_break_before(p):
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn('w:pageBreakBefore')) is None:
        pPr.insert(0, OxmlElement('w:pageBreakBefore'))


def set_table_borders(tb):
    tblPr = tb._tbl.tblPr
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


def fix_quotes_text(text):
    """把直引号/开闭错乱的中文引号按顺序重排成配对的“”。"""
    if not any(c in text for c in '"“”'):
        return text
    out, open_q = [], True
    for ch in text:
        if ch in '"“”':
            out.append('“' if open_q else '”')
            open_q = not open_q
        else:
            out.append(ch)
    return ''.join(out)


def apply(src, cfg, out=None):
    doc = Document(src)
    font = cfg['font']
    set_run = make_set_run(font)
    body = cfg['body']
    h_sizes = {st: info['size_pt'] for st, info in cfg['headings'].items()}

    # 1) 样式层兜底
    for sn in ('Normal', 'Body Text', 'First Paragraph', 'List Paragraph', 'Normal (Web)'):
        fix_style(doc, sn, font, body['size_pt'], None)
    for st, sz in h_sizes.items():
        fix_style(doc, st, font, sz, True)

    # 2) 段落层
    paras = doc.paragraphs
    for p in paras:
        st = p.style.name
        txt = p.text.strip()
        pf = p.paragraph_format
        if st in h_sizes:
            pf.line_spacing = body['line_spacing']
            pf.first_line_indent = None
            if cfg['chapter_page_break'] and st == 'Heading 1' and txt.startswith('第') and '章' in txt[:6]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_page_break_before(p)
            for r in p.runs:
                set_run(r, h_sizes[st], True)
        elif not st.startswith('toc'):
            pf.line_spacing = body['line_spacing']
            if txt:
                pf.first_line_indent = Emu(body['first_line_indent_emu'])
            for r in p.runs:
                set_run(r, body['size_pt'], r.bold if r.bold else None)

    # 3) 封面重做（首段是"xxx文件/报告/方案/公告"整行标题时）
    if cfg['cover'] and paras and paras[0].text.strip():
        first = paras[0]
        t0 = first.text.strip()
        doctype_kw = next((kw for kw in ('询比采购文件', '采购文件', '招标文件', '需求文件', '报告', '方案') if kw in t0), None)
        if doctype_kw and len(t0) > len(doctype_kw):
            head_part = t0[:t0.rindex(doctype_kw)]
            first.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = first._p.get_or_add_pPr()
            pb = pPr.find(qn('w:pageBreakBefore'))
            if pb is not None:
                pPr.remove(pb)
            for r in list(first.runs):
                r._r.getparent().remove(r._r)
            # 公司名与项目名拆行：公司名以"公司"结尾时拆开
            lines = []
            if '公司' in head_part:
                cut = head_part.index('公司') + 2
                lines = [head_part[:cut], head_part[cut:]]
            else:
                lines = [head_part]
            for line in lines:
                if line:
                    r = first.add_run(line)
                    set_run(r, cfg['cover_title_pt'], True)
                    r.add_break(WD_BREAK.LINE)
            r = first.add_run('')
            set_run(r, cfg['cover_title_pt'], True)
            r.add_break(WD_BREAK.LINE)
            r = first.add_run(doctype_kw)
            set_run(r, cfg['cover_doctype_pt'], True)
            first.paragraph_format.space_before = Pt(120)
            # 紧随其后的采购人/日期行居中
            for p in paras[1:3]:
                if p.style.name in h_sizes:
                    break
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = None
                for r in p.runs:
                    set_run(r, 15, True)

    # 4) 表格层
    for tb in doc.tables:
        set_table_borders(tb)
        for ri, row in enumerate(tb.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.first_line_indent = None
                    for r in p.runs:
                        set_run(r, cfg['table_cell_size_pt'], True if ri == 0 else None)

    # 5) 引号修复
    if cfg['fix_quotes']:
        def walk(paragraphs):
            for p in paragraphs:
                for r in p.runs:
                    if r.text and any(c in r.text for c in '"“”'):
                        r.text = fix_quotes_text(r.text)
        walk(doc.paragraphs)
        for tb in doc.tables:
            for row in tb.rows:
                for cell in row.cells:
                    walk(cell.paragraphs)

    dst = out or src
    doc.save(dst)
    return dst


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx_path')
    ap.add_argument('--config', help='格式配置 JSON（extract_fingerprint.py --json 产物或手写精简版）')
    ap.add_argument('--out', help='输出路径，缺省原地保存')
    args = ap.parse_args()
    cfg = load_config(args.config)
    dst = apply(args.docx_path, cfg, args.out)
    print('done:', dst)


if __name__ == '__main__':
    sys.exit(main())
